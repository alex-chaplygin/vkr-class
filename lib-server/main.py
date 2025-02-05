import json
import os.path
from sys import argv, exit
from docx import Document


def clean_text(text):
    '''
    Очистка от перевода строк
    '''
    return text.replace('\n', ' ')


def process(data):
    '''
    Замена строк
    '''
    code = data[10]
    if code != '09.04.04' and code != '09.03.04':
        print('Неправильная специальность:', data[1], code)
        exit(1)
    data[10] = code + ('.68' if code == '09.04.04' else '.62') + ' - Программная инженерия'
    data[11] = 'КАФЕДРА ПРОГРАММНОЙ ИНЖЕНЕРИИ'
    data[12] = '68 - Магистр' if code == '09.04.04' else '62 - Бакалавр'
    return data


def make_word(res):
    doc = Document()
    table = doc.add_table(rows=len(res), cols=4)
    for i in range(len(res)):
        cells = table.rows[i].cells
        cells[0].text = ''
        cells[1].text = res[i][1]
        cells[2].text = res[i][4]
        cells[3].text = res[i][0]
    doc.save('/tmp/act.docx')

if len(argv) != 3:
    print("main.py <prikaz> <pdf path>")
    exit(1)

document = Document(argv[1])
path = argv[2]
table = document.tables[1]
vkrs = []

keys = None
for i, row in enumerate(table.rows):
    text = (clean_text(cell.text) for cell in row.cells)
    if i == 0:
        continue
    vkrs.append(list(text))

res = []
for vkr in vkrs:
    vkr_1 = vkr[1].split(',')
    vkr_theme = vkr[2].lstrip().rstrip()
    vkr_name = vkr_1[0].lstrip().rstrip()
    vkr_id = vkr_1[1][6:17]
    name = path + vkr_id
    if not os.path.exists(name + ".pdf") or not os.path.exists(name + ".docx"):
        print("Нет ВКР и/или сведений", vkr_name, vkr_id)
    else:
        doc = Document(name + ".docx")
        tab = doc.tables[0]
        data = [clean_text(c.text).lstrip().rstrip() for c in tab.column_cells(1)][1:]
        if data[0] != vkr_id:
            print('Шифр не совпадает', vkr_id, data[0], sep='\n')
        if data[1] != vkr_name:
            print('ФИО не совпадает', vkr_name, data[1], sep='\n')
        if data[4] != vkr_theme:
            print('Тема не совпадает', vkr_id, vkr_name, vkr_theme, data[4], sep='\n')
        res.append(process(data))

make_word(res)

print(json.dumps(res))
